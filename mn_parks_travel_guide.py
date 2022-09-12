# Create MN State Parks travel guide in Word using API data

def get_data():
    # get_data function has no inputs
    # requests and stores data from API, return error message if applicable
    # returns list of dictionaries, park names and ID numbers
    import requests
    # get the API data
    url = "https://mn-state-parks.herokuapp.com/api/list"
    try:
        mn_parks = requests.get(url).json()
        # print('Data retrieved.')
        # mn_parks is a list of dictionaries, each of which has keys name and park_id, a 5-digit number
        return mn_parks
    except Exception:
        # https://www.30secondsofcode.org/articles/s/python-bare-except
        print("An error was encountered while requesting the data.")
        print("\n"*2)  # print empty lines
        return None


def choose_parks(mn_parks):
    # inputs list of dictionaries, park names and ID numbers
    # generates and returns list of 5 unique, randomly chosen parks
    import random
    my_parks = []  # empty list, will hold parkIDs
    count = 0   # will track number of unique parks selected
    park_id_list = []  # empty list
    for index in range(len(mn_parks)):
        park_id_list.append(mn_parks[index]['park_id'])  # create list of park IDs
    while count < 5:
        new_park = random.choice(park_id_list)  # randomly choose a park ID from the list
        # print("New park selected " + new_park)  # for debugging
        if new_park not in my_parks:
            count += 1  # if this park not already selected, add it to list and increment count
            my_parks.append(new_park)
    # print(my_parks)  # for debugging
    return my_parks


def write_word_doc(new_park, mn_parks_guide):
    # adds information to the Word doc for one park (new_park)
    import requests
    import docx
    import shutil  # not sure why the style guide prefers individual import statements
    # print("write_word_doc for ", new_park["name"])  # for debugging

    # add to Word doc: park name, address, and URL
    mn_parks_guide.add_heading(new_park["name"], 2)
    mn_parks_guide.add_heading("Contact Information:", 3)
    mn_parks_guide.add_heading("Address:", 5)
    mn_parks_guide.add_paragraph(new_park["address"], 'Normal')
    mn_parks_guide.add_heading("Website:", 5)
    mn_parks_guide.add_paragraph(new_park["url"], 'Normal')

    # the first park image will be added here, remaining after the rest of the text
    image_url = new_park["park_images"][0]
    image = requests.get(image_url, stream=True)  # download first image
    image_file_name = image_url.split("/")[-1]  # set file name
    # split the URL on the slashes and take the last segment for the file name
    # print("retrieved first image ", image_file_name)   # for debugging
    with open(image_file_name, 'wb') as f:  # write file
        shutil.copyfileobj(image.raw, f)  # https://www.scrapingbee.com/blog/download-image-python/
        # print("downloaded " + image_url)  # for debugging
        mn_parks_guide.add_picture(image_file_name, width=docx.shared.Inches(4))
    # https://towardsdatascience.com/how-to-download-an-image-using-python-38a75cfa21c
    # Pycharm is giving an error here, but it is working and when I tried other methods it did not work...
    # Pycharm error says "Cannot find reference 'shared' in "__init__.py"

    mn_parks_guide.add_heading("Highlights", 3)
    for item in new_park["highlights"]:
        # print(item)  # for debugging
        mn_parks_guide.add_paragraph(item, 'List Bullet 2')
        # bulleted list for the highlights

    # print header for park information section
    mn_parks_guide.add_heading("Interesting information about "+new_park["name"]+":", 2)
#  some parks are missing items from the park_information section, can cause key error
    if "Landscape" in new_park["park_information"]:
        mn_parks_guide.add_heading("Landscape", 4)
        mn_parks_guide.add_paragraph(new_park["park_information"]["Landscape"], 'Normal')
    if "Park Information" in new_park["park_information"]:
        mn_parks_guide.add_heading("Park Information", 4)
        mn_parks_guide.add_paragraph(new_park["park_information"]["Park Information"], 'Normal')
    if "History" in new_park["park_information"]:
        mn_parks_guide.add_heading("History", 4)
        mn_parks_guide.add_paragraph(new_park["park_information"]["History"], 'Normal')
    if "Geology" in new_park["park_information"]:
        mn_parks_guide.add_heading("Geology", 4)
        mn_parks_guide.add_paragraph(new_park["park_information"]["Geology"], 'Normal')
    if "Wildlife" in new_park["park_information"]:
        mn_parks_guide.add_heading("Wildlife", 4)
        mn_parks_guide.add_paragraph(new_park["park_information"]["Wildlife"], 'Normal')

    mn_parks_guide.save("MNParksGuide.docx")
    # print("Document saved.")  # for debugging

    # add images after park information
    # already used the first image, so can delete that from the list
    new_park["park_images"].pop(0)
    # print(new_park["park_images"])   # for debugging
    mn_parks_guide.add_heading("Photos from "+new_park["name"], 3)
    for image_url in new_park["park_images"]:   # for each image,
        # print(image_url)  # for debugging
        image = requests.get(image_url, stream=True)   # download the image
        image_file_name = image_url.split("/")[-1]    # set file name
        # split the URL on the slashes and take the last segment for the file name
        # print("retrieved image ", image_file_name)   # for debugging
        with open(image_file_name, 'wb') as f:        # write file
            shutil.copyfileobj(image.raw, f)  # https://www.scrapingbee.com/blog/download-image-python/
            # print("downloaded " + image_url)  # for debugging
            mn_parks_guide.add_picture(image_file_name, width=docx.shared.Inches(4))
        # https://towardsdatascience.com/how-to-download-an-image-using-python-38a75cfa21c
        # MNParksGuide.add_page_break()  # resulted in blank pages too often
        # https://www.geeksforgeeks.org/working-with-page-break-python-docx-module/
    mn_parks_guide.add_heading("-"*70, 2)  # print blank line
    mn_parks_guide.save("MNParksGuide.docx")


def parks_data(my_parks):
    # inputs list of specific parks
    # requests and stores data from API, returns error message if applicable
    # calls write_word_doc function
    import requests
    import docx
    mn_parks_guide = docx.Document()  # create new document
    mn_parks_guide.add_paragraph("Your Travel Guide for 5 Random Minnesota State Parks", 'Title')

    # get the API data
    urlbase = "https://mn-state-parks.herokuapp.com/api/"
    for park in my_parks:
        url = urlbase+park
        # print(url)  # for debugging

        try:
            new_park = requests.get(url).json()
            # print(park, ' Data retrieved.')  # for debugging
            # print(new_park)  # for debugging
        except Exception:
            print(park, " An error was encountered while requesting the data.")
            print("\n"*2)  # print empty lines
            new_park = None  # assign value so that return statement doesn't error out

        write_word_doc(new_park, mn_parks_guide)


def main():
    mn_parks = get_data()  # function to extract data from api
    if mn_parks is None:
        print("Please try again later.  An error occurred in retrieving the data.")
    else:
        print("Selecting parks...")
        my_parks = choose_parks(mn_parks)
        print("Retrieving parks data... this may take a minute or two.")
        parks_data(my_parks)
        print("Happy exploring!")


main()
